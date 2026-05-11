"""Visual crosshair calibration via a Word document.

MAKE (default):
    python tools/calib_docx.py --make -o out/calib.docx

    Generates a 2-page Word document:
      - The official W-3 form rendered behind the page.
      - One red floating text box per field / grid point, positioned at the
        current overlay coordinate.  Each box is labeled with the field name.
      - Drag a box to where the value should actually appear on the form.
      - Save the .docx (works in Word and Google Docs).

READ & APPLY:
    python tools/calib_docx.py out/calib_edited.docx --calib out/calib2.pdf
    python tools/calib_docx.py out/calib_edited.docx --patch

    Reads the dragged positions back, re-renders calib.pdf so you can verify
    visually, and/or patches src/plugfile/pdf_export.py constants in-place.

Coordinate conversion
---------------------
PDF user-space: origin bottom-left, y increases upward.
Word EMU:        origin top-left  of page, y increases downward.
1 pt = 12700 EMU; letter page = 612 x 792 pt = 7772160 x 10058400 EMU.

  x_emu = x_pdf_pt * 12700
  y_emu = (792 - y_pdf_pt) * 12700
  (reverse: x_pdf_pt = x_emu / 12700, y_pdf_pt = 792 - y_emu / 12700)
"""

from __future__ import annotations

import copy
import io
import os
import sys
import tempfile
from pathlib import Path
from typing import Any

# ---- path setup ------------------------------------------------------------
_HERE = Path(__file__).resolve().parent
_REPO = _HERE.parent
sys.path.insert(0, str(_REPO / "src"))

from plugfile.pdf_export import (
    CASING_COL_X,
    CASING_ROW_Y,
    PERF_PAIR_X,
    PERF_ROW_Y,
    PLUG_COL_X,
    PLUG_ROW_Y,
    W3_COORDS,
    FieldCoord,
)

# ---- imports ---------------------------------------------------------------
try:
    import pypdfium2 as pdfium
except ImportError:
    sys.exit("pypdfium2 required: pip install pypdfium2")

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Emu, Inches, RGBColor
    from docx.oxml import OxmlElement
    from lxml import etree
except ImportError:
    sys.exit("python-docx and lxml required: pip install python-docx lxml")

# ---- constants -------------------------------------------------------------
PAGE_W_PT = 612.0
PAGE_H_PT = 792.0
PAGE_W_EMU = int(PAGE_W_PT * 12700)
PAGE_H_EMU = int(PAGE_H_PT * 12700)

# Crosshair shape dimensions (pt)
SHAPE_W_PT = 100.0
SHAPE_H_PT = 14.0
SHAPE_W_EMU = int(SHAPE_W_PT * 12700)
SHAPE_H_EMU = int(SHAPE_H_PT * 12700)

# Half-dimensions for centering the shape on the target coordinate
HALF_W_PT = SHAPE_W_PT / 2
HALF_H_PT = SHAPE_H_PT / 2


# ---- coordinate conversion -------------------------------------------------

def pdf_to_word_emu(x_pt: float, y_pt: float) -> tuple[int, int]:
    """Shape anchor (top-left) in Word EMU, centered on the PDF coordinate."""
    x_emu = int((x_pt - HALF_W_PT) * 12700)
    y_emu = int((PAGE_H_PT - y_pt - HALF_H_PT) * 12700)
    return max(0, x_emu), max(0, y_emu)


def word_emu_to_pdf(x_emu: int, y_emu: int) -> tuple[float, float]:
    """Reverse: Word anchor EMU -> PDF coordinate (center of shape)."""
    x_pt = x_emu / 12700 + HALF_W_PT
    y_pt = PAGE_H_PT - (y_emu / 12700) - HALF_H_PT
    return round(x_pt, 1), round(y_pt, 1)


# ---- shape name encoding ---------------------------------------------------
# Names encode shape type + key so read-back knows which constant to update.
# Format: "wp:<type>:<key>"  (no spaces, colon-delimited)
# Types: "scalar", "plug_col", "plug_row", "cas_row", "cas_col",
#        "perf_row", "perf_pair"

def _name(kind: str, key: str) -> str:
    return f"wp:{kind}:{key}"


def _parse_name(name: str) -> tuple[str, str] | None:
    parts = name.split(":", 2)
    if len(parts) == 3 and parts[0] == "wp":
        return parts[1], parts[2]
    return None


# ---- collect all shapes to place -------------------------------------------

def _all_shapes(page: int) -> list[tuple[str, float, float]]:
    """Return list of (shape_name, x_pt, y_pt) for all crosshairs on `page`."""
    shapes: list[tuple[str, float, float]] = []

    # Scalar fields
    for name, coord in W3_COORDS.items():
        if coord.page == page:
            shapes.append((_name("scalar", name), coord.x, coord.y))

    if page == 0:
        # Plug column reference line (at cement_date row)
        y_ref = PLUG_ROW_Y.get("cement_date", 460.0)
        for i, x in enumerate(PLUG_COL_X, 1):
            shapes.append((_name("plug_col", str(i)), x, y_ref))

        # Plug row reference line (at column 1 x)
        x_ref = PLUG_COL_X[0] if PLUG_COL_X else 220.0
        for row_name, y in PLUG_ROW_Y.items():
            shapes.append((_name("plug_row", row_name), x_ref, y))

        # Casing rows (SIZE column x)
        x_cas = CASING_COL_X.get("size", 40.0)
        for i, y in enumerate(CASING_ROW_Y, 1):
            shapes.append((_name("cas_row", str(i)), x_cas, y))

        # Casing cols (first row y)
        y_cas = CASING_ROW_Y[0] if CASING_ROW_Y else 325.0
        for col_name, x in CASING_COL_X.items():
            shapes.append((_name("cas_col", col_name), x, y_cas))

        # Perf rows (left pair from_x)
        x_perf = PERF_PAIR_X[0][0] if PERF_PAIR_X else 75.0
        for i, y in enumerate(PERF_ROW_Y, 1):
            shapes.append((_name("perf_row", str(i)), x_perf, y))

        # Perf pairs (first row y)
        y_perf = PERF_ROW_Y[0] if PERF_ROW_Y else 251.0
        for i, (from_x, to_x) in enumerate(PERF_PAIR_X, 1):
            shapes.append((_name("perf_pair", f"{i}:from"), from_x, y_perf))
            shapes.append((_name("perf_pair", f"{i}:to"), to_x, y_perf))

    return shapes


# ---- Word XML helpers ------------------------------------------------------

_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _make_anchor_xml(
    shape_id: int,
    name: str,
    x_emu: int,
    y_emu: int,
    width_emu: int,
    height_emu: int,
    behind: bool = False,
) -> str:
    """Return the `<w:drawing>` XML string for one floating shape."""
    z = "1" if behind else str(251658240 + shape_id)
    label = name.split(":", 2)[-1] if ":" in name else name  # human-readable label

    # inline color: use gray for background image, red for crosshairs
    ln_color = "CCCCCC" if behind else "FF0000"
    txt_color = "FF0000"

    return f"""<w:drawing
    xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
    xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
    xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
    xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"
    xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"
    xmlns:o="urn:schemas-microsoft-com:office:office"
    xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
    xmlns:v="urn:schemas-microsoft-com:vml"
    xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:w10="urn:schemas-microsoft-com:office:word"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
    xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
    xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"
    xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
    xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"
    xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"
    xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
    xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
    xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
    xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
    xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
             relativeHeight="{z}" behindDoc="{'1' if behind else '0'}"
             locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{x_emu}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="page">
      <wp:posOffset>{y_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{shape_id}" name="{name}"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr txBx="1">
            <a:spLocks noChangeArrowheads="1"/>
          </wps:cNvSpPr>
          <wps:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:noFill/>
            <a:ln w="19050">
              <a:solidFill><a:srgbClr val="{ln_color}"/></a:solidFill>
            </a:ln>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent>
              <w:p>
                <w:pPr>
                  <w:jc w:val="center"/>
                  <w:spacing w:before="0" w:after="0"/>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:color w:val="{txt_color}"/>
                    <w:sz w:val="14"/>
                    <w:szCs w:val="14"/>
                  </w:rPr>
                  <w:t>+ {label}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr lIns="36000" rIns="36000" tIns="36000" bIns="36000"
                      anchor="ctr" anchorCtr="1">
            <a:noAutofit/>
          </wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>"""


def _make_image_anchor_xml(
    shape_id: int, rel_id: str,
    x_emu: int, y_emu: int,
    width_emu: int, height_emu: int,
) -> str:
    """Return the `<w:drawing>` XML for a floating behind-text picture."""
    return f"""<w:drawing
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
             relativeHeight="1" behindDoc="1"
             locked="1" layoutInCell="1" allowOverlap="0">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page">
      <wp:posOffset>{x_emu}</wp:posOffset>
    </wp:positionH>
    <wp:positionV relativeFrom="page">
      <wp:posOffset>{y_emu}</wp:posOffset>
    </wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{shape_id}" name="bg{shape_id}"/>
    <wp:cNvGraphicFramePr>
      <a:graphicFrameLocks noChangeAspect="1"/>
    </wp:cNvGraphicFramePr>
    <a:graphic>
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic>
          <pic:nvPicPr>
            <pic:cNvPr id="{shape_id}" name="bg{shape_id}"/>
            <pic:cNvPicPr>
              <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
            </pic:cNvPicPr>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{rel_id}"/>
            <a:stretch><a:fillRect/></a:stretch>
          </pic:blipFill>
          <pic:spPr bwMode="auto">
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>"""


def _xml_drawing_to_run(drawing_xml: str, doc: "Document") -> Any:
    """Parse drawing XML and wrap it in a <w:r> element."""
    drawing_el = etree.fromstring(drawing_xml)
    run = OxmlElement("w:r")
    run.append(drawing_el)
    return run


# ---- MAKE ------------------------------------------------------------------

def make_docx(dest: Path, template_path: Path | None = None) -> None:
    """Generate the calibration Word document."""
    from plugfile.pdf_export import _resolve_template
    template_path = _resolve_template(template_path)

    # Render form pages to in-memory PNGs
    pdf_doc = pdfium.PdfDocument(str(template_path))
    png_buffers: list[io.BytesIO] = []
    for i in range(min(2, len(pdf_doc))):
        bm = pdf_doc[i].render(scale=2.0)
        pil = bm.to_pil()
        buf = io.BytesIO()
        pil.save(buf, format="PNG")
        buf.seek(0)
        png_buffers.append(buf)

    doc = Document()

    # Zero margins, letter page
    for section in doc.sections:
        section.page_width  = Pt(PAGE_W_PT)
        section.page_height = Pt(PAGE_H_PT)
        section.left_margin   = Pt(0)
        section.right_margin  = Pt(0)
        section.top_margin    = Pt(0)
        section.bottom_margin = Pt(0)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)

    shape_id = 100  # start above Word's default IDs
    _A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    _R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for page_idx, png_buf in enumerate(png_buffers):
        # One paragraph per page with all shapes attached as runs
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # Use add_picture (inline) to correctly register the image part and
        # get the relationship ID, then replace the inline drawing with our
        # floating behind-text anchor.
        probe_run = para.add_run()
        probe_run.add_picture(png_buf, width=Pt(PAGE_W_PT))
        probe_drawing = probe_run._element.find(qn("w:drawing"))
        blip_tag = f"{{{_A_NS}}}blip"
        blip_el = next(probe_drawing.iter(blip_tag), None)
        if blip_el is None:
            raise RuntimeError("Could not find blip in inline picture element")
        rel = blip_el.get(f"{{{_R_NS}}}embed")
        probe_run._element.remove(probe_drawing)
        # Remove the now-empty probe run
        para._element.remove(probe_run._element)

        # Background image run (behind-text, full page)
        bg_xml = _make_image_anchor_xml(
            shape_id, rel, 0, 0, PAGE_W_EMU, PAGE_H_EMU
        )
        para._element.append(_xml_drawing_to_run(bg_xml, doc))
        shape_id += 1

        # Crosshair shapes
        for shp_name, x_pt, y_pt in _all_shapes(page_idx):
            x_emu, y_emu = pdf_to_word_emu(x_pt, y_pt)
            anchor_xml = _make_anchor_xml(
                shape_id, shp_name, x_emu, y_emu,
                SHAPE_W_EMU, SHAPE_H_EMU, behind=False,
            )
            para._element.append(_xml_drawing_to_run(anchor_xml, doc))
            shape_id += 1

        # Page break (except after the last page)
        if page_idx < len(png_buffers) - 1:
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run_el = OxmlElement("w:r")
            run_el.append(br)
            para._element.append(run_el)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))

    total_shapes = sum(len(_all_shapes(i)) for i in range(len(png_buffers)))
    print(f"Wrote {dest}  ({total_shapes} crosshairs across {len(png_buffers)} pages)")
    print()
    print("Next steps:")
    print("  1. Open the .docx in Word or Google Docs.")
    print("  2. Drag each red labeled box to where that field's value belongs.")
    print("  3. Save/download as .docx.")
    print("  4. python tools/calib_docx.py <edited.docx> --calib out/calib2.pdf")
    print("     Open calib2.pdf to verify the moved crosshairs on the actual form.")
    print("  5. When satisfied: python tools/calib_docx.py <edited.docx> --patch")


# ---- READ ------------------------------------------------------------------

def read_docx(src: Path) -> dict[str, tuple[float, float]]:
    """Extract shape positions from the edited calibration docx.

    Returns a dict mapping shape names (e.g. 'wp:scalar:api_number') to
    (x_pdf_pt, y_pdf_pt) tuples."""
    doc = Document(str(src))
    results: dict[str, tuple[float, float]] = {}

    # Walk all drawing elements in the document body
    ns = {"wp": _WP, "w": _W}
    for drawing in doc.element.body.iter(qn("w:drawing")):
        for anchor in drawing.iter(qn("wp:anchor")):
            docpr = anchor.find(qn("wp:docPr"))
            if docpr is None:
                continue
            name = docpr.get("name", "")
            if not name.startswith("wp:"):
                continue

            # Extract positionH and positionV offsets
            ph = anchor.find(f".//{qn('wp:positionH')}/{qn('wp:posOffset')}")
            pv = anchor.find(f".//{qn('wp:positionV')}/{qn('wp:posOffset')}")
            if ph is None or pv is None:
                continue
            x_emu = int(ph.text or 0)
            y_emu = int(pv.text or 0)
            x_pt, y_pt = word_emu_to_pdf(x_emu, y_emu)
            results[name] = (x_pt, y_pt)

    return results


def positions_to_constants(
    positions: dict[str, tuple[float, float]],
) -> dict:
    """Convert read-back positions into the constants expected by pdf_export."""
    coords = dict(W3_COORDS)
    plug_col_x = list(PLUG_COL_X)
    plug_row_y = dict(PLUG_ROW_Y)
    casing_row_y = list(CASING_ROW_Y)
    casing_col_x = dict(CASING_COL_X)
    perf_row_y = list(PERF_ROW_Y)
    perf_pair_x = [list(pair) for pair in PERF_PAIR_X]

    # plug_col reference row (cement_date) — only x matters for columns
    cement_date_y = plug_row_y.get("cement_date", 460.0)

    for name, (x_pt, y_pt) in positions.items():
        parsed = _parse_name(name)
        if parsed is None:
            continue
        kind, key = parsed

        if kind == "scalar" and key in coords:
            old = coords[key]
            coords[key] = FieldCoord(
                page=old.page, x=x_pt, y=y_pt,
                max_width=old.max_width, font_size=old.font_size,
            )
        elif kind == "plug_col":
            idx = int(key) - 1
            if 0 <= idx < len(plug_col_x):
                plug_col_x[idx] = x_pt  # x defines column position
        elif kind == "plug_row":
            if key in plug_row_y:
                plug_row_y[key] = y_pt  # y defines row position
        elif kind == "cas_row":
            idx = int(key) - 1
            if 0 <= idx < len(casing_row_y):
                casing_row_y[idx] = y_pt
        elif kind == "cas_col":
            if key in casing_col_x:
                casing_col_x[key] = x_pt
        elif kind == "perf_row":
            idx = int(key) - 1
            if 0 <= idx < len(perf_row_y):
                perf_row_y[idx] = y_pt
        elif kind == "perf_pair":
            parts = key.split(":")
            if len(parts) == 2:
                pair_idx = int(parts[0]) - 1
                from_or_to = parts[1]
                if 0 <= pair_idx < len(perf_pair_x):
                    slot = 0 if from_or_to == "from" else 1
                    perf_pair_x[pair_idx][slot] = x_pt

    return {
        "coords": coords,
        "plug_col_x": tuple(plug_col_x),
        "plug_row_y": plug_row_y,
        "casing_row_y": tuple(casing_row_y),
        "casing_col_x": casing_col_x,
        "perf_row_y": tuple(perf_row_y),
        "perf_pair_x": tuple(tuple(p) for p in perf_pair_x),
    }


# ---- CLI -------------------------------------------------------------------

def _cli():
    import argparse
    p = argparse.ArgumentParser(
        prog="calib_docx",
        description="Visual W-3 crosshair calibration via Word document.",
    )
    p.add_argument("docx", nargs="?",
                   help="Edited calibration .docx (omit when using --make)")
    p.add_argument("--make", action="store_true",
                   help="Generate the calibration docx (default when no docx given)")
    p.add_argument("-o", "--output", default="out/calib.docx",
                   help="Output path for --make (default: out/calib.docx)")
    p.add_argument("--template", default=None,
                   help="Override path to w-3p.pdf")
    p.add_argument("--calib", metavar="PDF",
                   help="After reading, re-render calib overlay to this PDF path")
    p.add_argument("--patch", action="store_true",
                   help="Patch src/plugfile/pdf_export.py constants in-place")
    a = p.parse_args()

    if a.make or a.docx is None:
        dest = Path(a.output)
        make_docx(dest, Path(a.template) if a.template else None)
        return

    src = Path(a.docx)
    if not src.exists():
        sys.exit(f"File not found: {src}")

    print(f"Reading {src} ...")
    positions = read_docx(src)
    print(f"  {len(positions)} crosshair positions extracted.")
    consts = positions_to_constants(positions)

    if a.calib:
        import plugfile.pdf_export as _mod
        _mod.W3_COORDS    = consts["coords"]
        _mod.PLUG_COL_X   = consts["plug_col_x"]
        _mod.PLUG_ROW_Y   = consts["plug_row_y"]
        _mod.CASING_ROW_Y = consts["casing_row_y"]
        _mod.CASING_COL_X = consts["casing_col_x"]
        _mod.PERF_ROW_Y   = consts["perf_row_y"]
        _mod.PERF_PAIR_X  = consts["perf_pair_x"]
        calib_dest = Path(a.calib)
        calib_dest.parent.mkdir(parents=True, exist_ok=True)
        pdf = _mod.render_calibration_overlay()
        calib_dest.write_bytes(pdf)
        print(f"Calibration overlay -> {calib_dest} ({len(pdf):,} bytes)")

    if a.patch:
        # Reuse apply_coords patch logic
        sys.path.insert(0, str(_HERE))
        from apply_coords import patch_source
        patch_source(
            consts["coords"],
            consts["plug_col_x"],
            consts["plug_row_y"],
            consts["casing_row_y"],
            consts["casing_col_x"],
            consts["perf_row_y"],
            consts["perf_pair_x"],
        )
        print("Patched pdf_export.py -- run pytest to verify.")


if __name__ == "__main__":
    _cli()
